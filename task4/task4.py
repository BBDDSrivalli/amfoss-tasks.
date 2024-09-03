import os
import csv
import requests
import logging
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler, CallbackQueryHandler,
    ConversationHandler, ContextTypes, filters
)
from dotenv import load_dotenv
from docx import Document

# Load environment variables
load_dotenv()

# Constants for API keys
TELEGRAM_TOKEN = os.getenv('TELEGRAM_TOKEN', '7532229434:AAHO1RK7iNlO9dxlKNhYKnu7V_sBAiK4_-U')
BOOKS_API_KEY = os.getenv('BOOKS_API_KEY', 'AIzaSyBx27rgeFfi_5ToJwqJRCGGUwh5T_X-__I')

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Genres List
GENRES = [
    "Adventure", "Fantasy", "Mystery", "Romance", "Science Fiction",
    "Thriller", "Non-fiction", "Historical", "Horror", "Biography"
]

# User-specific reading list
user_bookshelf = {}

# Conversation states
CHOOSE_GENRE, GET_BOOK_TITLE, HANDLE_SHELF = range(3)

# Start command
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    welcome_message = (
        "Hi there, I'm Page Pal, your friendly bookworm buddy! \n"
        "I’m here to help you find awesome books and keep track of your reading adventures!\n"
        "You can start by typing /browse to explore book genres or /myshelf to manage your personal bookshelf."
    )
    await update.message.reply_text(welcome_message)

# Browse genres
async def browse_genres(update: Update, context: ContextTypes.DEFAULT_TYPE):
    genre_buttons = [
        [InlineKeyboardButton(genre, callback_data=genre.lower())] for genre in GENRES
    ]
    reply_markup = InlineKeyboardMarkup(genre_buttons)
    await update.message.reply_text(
        "Yay! Let's find some books! \nPick a genre:",
        reply_markup=reply_markup
    )
    return CHOOSE_GENRE

# Show books for a selected genre
async def show_books(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    genre = query.data
    books = fetch_books_by_genre(genre)
    
    if books:
        filename = f"{genre}_books.csv"
        try:
            with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(["Title", "Authors", "Description", "Published Date", "Preview Link"])
                writer.writerows(books)
            
            with open(filename, 'rb') as file:
                await query.message.reply_document(document=file)
        except Exception as e:
            logger.error(f"Error writing or sending CSV file: {e}")
            await query.message.reply_text("There was an error creating the book list. Please try again later.")
        finally:
            if os.path.exists(filename):
                os.remove(filename)
        
    else:
        await query.message.reply_text("Oops! I couldn't find any books in that genre.")
    
    await query.message.reply_text("Want to explore another genre? Just type /browse again!")
    return ConversationHandler.END

# Fetch books using Google Books API
def fetch_books_by_genre(genre):
    try:
        url = f"https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={BOOKS_API_KEY}"
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()
    except requests.exceptions.RequestException as e:
        logger.error(f"Error fetching books: {e}")
        return []

    book_list = []
    for item in data.get("items", []):
        info = item.get("volumeInfo", {})
        title = info.get("title", "Unknown")
        authors = ", ".join(info.get("authors", []))
        description = info.get("description", "No description available.")
        published_date = info.get("publishedDate", "N/A")
        preview_link = info.get("previewLink", "No preview available.")
        book_list.append([title, authors, description, published_date, preview_link])

    return book_list

# Manage personal bookshelf
async def my_shelf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    if user_id not in user_bookshelf:
        user_bookshelf[user_id] = []

    buttons = [
        [InlineKeyboardButton("Add a Book", callback_data='add')],
        [InlineKeyboardButton("Remove a Book ", callback_data='remove')],
        [InlineKeyboardButton("View My Bookshelf ", callback_data='view')]
    ]
    reply_markup = InlineKeyboardMarkup(buttons)
    await update.message.reply_text(
        "What would you like to do with your bookshelf today?",
        reply_markup=reply_markup
    )
    return HANDLE_SHELF

# Handle bookshelf actions
async def handle_shelf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    user_id = query.from_user.id

    if query.data == 'add':
        await query.message.reply_text("What’s the title of the book you want to add?")
        context.user_data['action'] = 'add'
        return GET_BOOK_TITLE

    elif query.data == 'remove':
        if not user_bookshelf.get(user_id):
            await query.message.reply_text("Your bookshelf is empty! Nothing to remove.")
            return ConversationHandler.END

        await query.message.reply_text("Which book would you like to remove?")
        context.user_data['action'] = 'remove'
        return GET_BOOK_TITLE

    elif query.data == 'view':
        if not user_bookshelf.get(user_id):
            await query.message.reply_text("Your bookshelf is empty! Start adding some books! 📚")
        else:
            doc = Document()
            doc.add_heading('My Cute Little Bookshelf', level=1)

            for book in user_bookshelf[user_id]:
                link = fetch_preview_link(book)
                doc.add_paragraph(f"{book} - Preview: {link or 'No preview available'}")

            doc_filename = 'bookshelf.docx'
            try:
                doc.save(doc_filename)
                with open(doc_filename, 'rb') as file:
                    await query.message.reply_document(document=file)
            except Exception as e:
                logger.error(f"Error creating or sending DOCX file: {e}")
                await query.message.reply_text("There was an error creating your bookshelf document. Please try again later.")
            finally:
                if os.path.exists(doc_filename):
                    os.remove(doc_filename)

        return ConversationHandler.END

# Add or remove books from the bookshelf
async def modify_shelf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    book_title = update.message.text

    if context.user_data.get('action') == 'add':
        if book_title not in user_bookshelf.get(user_id, []):
            user_bookshelf[user_id].append(book_title)
            await update.message.reply_text(f"'{book_title}' is now snugly placed on your bookshelf! 📚")
        else:
            await update.message.reply_text(f"'{book_title}' is already on your bookshelf.")
    elif context.user_data.get('action') == 'remove':
        if book_title in user_bookshelf.get(user_id, []):
            user_bookshelf[user_id].remove(book_title)
            await update.message.reply_text(f"'{book_title}' has been taken off your bookshelf! 🗑️")
        else:
            await update.message.reply_text(f"Couldn't find '{book_title}' on your bookshelf. 🤔")
    return ConversationHandler.END

# Fetch preview link for a specific book
def fetch_preview_link(book_title):
    try:
        url = f"https://www.googleapis.com/books/v1/volumes?q={book_title}&key={BOOKS_API_KEY}"
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()
    except requests.exceptions.RequestException as e:
        logger.error(f"Error fetching preview link: {e}")
        return None

    items = data.get("items", [])
    if items:
        return items[0].get("volumeInfo", {}).get("previewLink")
    return None

# Run the bot
def main():
    app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()

    # Conversation handler for genre browsing and bookshelf management
    conv_handler = ConversationHandler(
        entry_points=[
            CommandHandler('browse', browse_genres),
            CommandHandler('myshelf', my_shelf),
        ],
        states={
            CHOOSE_GENRE: [CallbackQueryHandler(show_books)],
            GET_BOOK_TITLE: [MessageHandler(filters.TEXT & ~filters.COMMAND, modify_shelf)],
            HANDLE_SHELF: [CallbackQueryHandler(handle_shelf)],
        },
        fallbacks=[CommandHandler('start', start)]
    )

    # Register handlers
    app.add_handler(CommandHandler('start', start))
    app.add_handler(conv_handler)

    logger.info("Bookie is up and ready to recommend books!")
    app.run_polling()

if __name__ == "__main__":
    main()

